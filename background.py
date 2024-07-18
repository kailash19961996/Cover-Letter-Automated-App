import base64
import time
import streamlit as st
from io import BytesIO
import openai
import os
import fitz
import requests
from docx import Document
from bs4 import BeautifulSoup

models = ["gpt-3.5-turbo", "gpt-4o"]
chosen_model = models[0]

def add_bg_from_local(image_file):
    with open(image_file, "rb") as image_file:
        encoded_string = base64.b64encode(image_file.read())
    st.markdown(
    f"""
    <style>
    .stApp {{
        background-image: url(data:image/{"jpg"};base64,{encoded_string.decode()});
        background-size: cover
    }}
    </style>
    """,
    unsafe_allow_html=True
    )

def show_gif_overlay(gif_path, duration):
    gif_html = f"""
        <style>
        .gif-overlay {{
            position: fixed;
            top: 75%;
            left: 50%;
            transform: translate(-50%, -50%);
            z-index: 9999;
        }}
        </style>
        <div class="gif-overlay">
            <img src="data:image/gif;base64,{base64.b64encode(open(gif_path, "rb").read()).decode()}" alt="Overlay GIF">
        </div>
    """
    overlay_placeholder = st.empty()
    overlay_placeholder.markdown(gif_html, unsafe_allow_html=True)
    time.sleep(duration)
    overlay_placeholder.empty()

def analyze_match(resume, job_details):
    prompt = f"Tell me in less than 150 words, Is the resume: \"{resume}\" and the job details: \"{job_details}\" relevant?. If and only relevant, give suggestions on missing skills and improvements. If irrelevant, let me know that this job posting is irrelevant"
    response = openai.ChatCompletion.create(
        model = chosen_model,
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=300
    )
    return response.choices[0].message['content']

cl_format = """
Kailash Subramaniyam \n
kylasben@gmail.com \n
[date]

Job Poster or Hiring manager name (If Job poster or hiring manager not available in job_details, just the company name from job_details) \n
Company name (search for the company name in the job_details)

Dear Hiring Manager,
Express why I want to work for this company and this role and how my experience is of relevance to this role

Sincerely,
Kailash Kumar
"""

def create_cv(resume, job_details):
    prompt = f"In less than 200 words, Generate a cover letter based on the resume: \"{resume}\" and the job details: \"{job_details}\" in this format: \"{cl_format}\":"
    response = openai.ChatCompletion.create(
        model = chosen_model,
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=400
    )
    return response.choices[0].message['content']

def add_adjustment(resume, job_details, new_cv, adjustments):
    prompt = f"In less than 200 words, Modify this cover letter: \"{new_cv}\", with only this adjustments: \"{adjustments}\", wherever applicable. Have a look at the resume: \"{resume}\" and job details: \"{job_details}\" for your reference."
    response = openai.ChatCompletion.create(
        model= "gpt-4o",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=400
    )
    return response.choices[0].message['content']


def extract_text_from_pdf(file):
    text = ""
    pdf_document = fitz.open(stream=file.read(), filetype="pdf")
    for page in pdf_document:
        text += page.get_text()
    return text

def extract_text_from_docx(file):
    doc = Document(file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def scrape_website(url):
    try:
        response = requests.get(url)
        response.raise_for_status()  # Raise HTTPError for bad responses
        soup = BeautifulSoup(response.content, 'html.parser')
        return soup
    except requests.exceptions.RequestException as e:
        st.error(f"An error occurred while trying to retrieve the data: {e}")
        return None

def create_docx(content):
    doc = Document()
    doc.add_paragraph(content)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


